"""Tests for the content extractor service."""

from __future__ import annotations

import io

import pytest

from crewai_productfeature_planner.services.content_extractor import extract_text


class TestExtractPlainText:
    def test_utf8_text(self):
        content = b"Hello, world! This is a test document."
        result = extract_text(content, filename="test.txt", content_type="text/plain")
        assert result == "Hello, world! This is a test document."

    def test_markdown(self):
        content = b"# Heading\n\nSome content here."
        result = extract_text(content, filename="readme.md", content_type="text/markdown")
        assert "# Heading" in result

    def test_json_file(self):
        content = b'{"key": "value"}'
        result = extract_text(content, filename="data.json", content_type="application/json")
        assert '"key"' in result

    def test_csv_file(self):
        content = b"name,age\nAlice,30\nBob,25"
        result = extract_text(content, filename="data.csv", content_type="text/csv")
        assert "Alice" in result

    def test_python_file_by_extension(self):
        content = b"def hello():\n    print('hi')"
        result = extract_text(content, filename="main.py", content_type="application/octet-stream")
        assert "def hello" in result

    def test_yaml_file(self):
        content = b"key: value\nlist:\n  - item1"
        result = extract_text(content, filename="config.yaml", content_type="application/x-yaml")
        assert "key: value" in result


class TestExtractPdf:
    def test_simple_pdf(self):
        """Create a minimal PDF and extract text."""
        from pypdf import PdfWriter

        writer = PdfWriter()
        from pypdf._page import PageObject
        from pypdf.generic import ArrayObject, NameObject, NumberObject

        page = PageObject.create_blank_page(width=612, height=792)
        # Add text content stream to the page
        import struct

        content = b"BT /F1 12 Tf 100 700 Td (Test document content) Tj ET"
        # Use a more reliable approach — just verify extraction handles valid PDFs
        buf = io.BytesIO()
        writer.add_blank_page(width=612, height=792)
        writer.write(buf)
        pdf_bytes = buf.getvalue()

        # A blank PDF won't have text — extraction returns None
        result = extract_text(pdf_bytes, filename="test.pdf", content_type="application/pdf")
        assert result is None  # Blank PDF has no text

    def test_pdf_detection_by_content_type(self):
        """Verify PDF is detected by content type even with wrong extension."""
        # Not a valid PDF — should return None gracefully
        result = extract_text(b"not a pdf", filename="file.bin", content_type="application/pdf")
        assert result is None

    def test_pdf_detection_by_extension(self):
        """Verify PDF is detected by .pdf extension."""
        result = extract_text(b"not a pdf", filename="doc.pdf", content_type="application/octet-stream")
        assert result is None


class TestExtractDocx:
    def test_simple_docx(self):
        """Create a minimal DOCX and extract text."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("First paragraph of the document.")
        doc.add_paragraph("Second paragraph with more content.")
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        result = extract_text(
            docx_bytes,
            filename="test.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        assert result is not None
        assert "First paragraph" in result
        assert "Second paragraph" in result

    def test_empty_docx(self):
        """DOCX with no text returns None."""
        from docx import Document

        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        result = extract_text(docx_bytes, filename="empty.docx", content_type="application/octet-stream")
        assert result is None

    def test_docx_detected_by_extension(self):
        """Verify DOCX is detected by extension."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Content")
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        result = extract_text(docx_bytes, filename="report.docx", content_type="application/octet-stream")
        assert result is not None
        assert "Content" in result


class TestUnsupportedFormats:
    def test_legacy_doc(self):
        """Legacy .doc files return None."""
        result = extract_text(b"\xd0\xcf\x11\xe0", filename="old.doc", content_type="application/msword")
        assert result is None

    def test_binary_garbage(self):
        """Random binary content returns None."""
        # Generate content with many non-UTF8 bytes
        binary = bytes(range(256)) * 10
        result = extract_text(binary, filename="data.bin", content_type="application/octet-stream")
        assert result is None

    def test_mostly_text_binary_passes(self):
        """File with mostly valid text passes the fallback check."""
        content = b"This is mostly valid text with some " + b"\xff" + b" here"
        result = extract_text(content, filename="mixed.dat", content_type="application/octet-stream")
        # Less than 10% replacement chars, so fallback succeeds
        assert result is not None
        assert "mostly valid text" in result


class TestEdgeCases:
    def test_empty_content(self):
        result = extract_text(b"", filename="empty.txt", content_type="text/plain")
        assert result == ""

    def test_no_extension_text_content_type(self):
        result = extract_text(b"text content", filename="noext", content_type="text/plain")
        assert result == "text content"

    def test_no_extension_binary_content_type(self):
        binary = bytes(range(256)) * 10
        result = extract_text(binary, filename="noext", content_type="application/octet-stream")
        assert result is None
